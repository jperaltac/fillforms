#!/usr/bin/env python3
"""
Simple generator of ANID Form 1 & Form 2 (2026) per student, from a CSV.

Requirements:
  - Python 3.9+
  - pip install python-docx pandas

Usage:
  1) Put your CSV next to this script (or pass --csv path).
  2) Keep the original templates with their current names in the same folder or pass --f1/--f2.
  3) Run:
        python generate_forms.py --csv estudiantes.csv
     The script will create ./salida/<Apellido_Nombre>_F1.docx and _F2.docx for each row.

Notes:
  - This script modifies text inline for each label it can match robustly.
  - For Formulario N°2 "checkboxes", we emulate checks by adding "[X]" before the chosen status.
  - If a label is not found exactly (templates change), the script logs a warning but continues.
"""

import argparse
import sys
from pathlib import Path
import pandas as pd
from docx import Document

# -------- Utilities --------
def full_name(row):
    nombre = str(row.get("nombre", "")).strip()
    apellido = str(row.get("apellido", "")).strip()
    return f"{nombre} {apellido}".strip()

def safe_str(val):
    return "" if pd.isna(val) else str(val)

def replace_line_startswith(doc, label_map):
    """
    Replace any paragraph whose text starts with a given label with "label value".
    Returns a set of labels that were successfully replaced.
    """
    done = set()
    # Start with normal paragraphs
    paragraphs = list(doc.paragraphs)
    # Also inspect tables:
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for p in paragraphs:
        text = p.text.strip()
        for label, value in label_map.items():
            if label in done:
                continue
            if text.startswith(label):
                # Write "label value" on the same line
                new_text = f"{label} {value}".rstrip()
                # Clear existing runs, then add one
                for r in p.runs:
                    r.clear()
                p.add_run(new_text)
                done.add(label)
    return done

def mark_status_checkboxes(doc, estado):
    """
    For Form 2: emulate a checkbox by prefixing the chosen line with "[X]" and others with "[ ]".
    Expected estado values (lowercase): postulacion_formal | aceptado | alumno_regular
    """
    estado = str(estado or "").strip().lower()
    targets = {
        "postulacion_formal": "En proceso de postulación formal",
        "aceptado": "Aceptado/a",
        "alumno_regular": "En calidad de Alumno/a Regular",
    }
    chosen_label = None
    if estado in targets:
        chosen_label = targets[estado]
    else:
        # try a few fallbacks
        if "postula" in estado:
            chosen_label = targets["postulacion_formal"]
        elif "regular" in estado:
            chosen_label = targets["alumno_regular"]
        elif "acept" in estado:
            chosen_label = targets["aceptado"]

    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for p in paragraphs:
        raw = p.text.strip()
        for lbl in targets.values():
            if raw.startswith(lbl):
                mark = "[X]" if lbl == chosen_label else "[ ]"
                new_text = f"{mark} {lbl}"
                for r in p.runs:
                    r.clear()
                p.add_run(new_text)

# -------- Fillers --------
def fill_form_1(doc, row):
    """
    Formulario N° 1: Certificado de Nota y Ranking para estudios de Pregrado
    """
    nombre_completo = full_name(row)
    label_map = {
        "Nombre del/de la Estudiante:": nombre_completo,
        "RUN o número de pasaporte:": safe_str(row.get("run") or row.get("pasaporte")),
        "Universidad de Pregrado*:": safe_str(row.get("universidad_pregrado")),
        "Programa de Estudios de pregrado*:": safe_str(row.get("programa_pregrado")),
        "Número de semestres de duración del programa académico de pregrado:": safe_str(row.get("semestres_pregrado")),
        "Región donde cursó los estudios de pregrado:": safe_str(row.get("region_pregrado")),
        # Notas
        "PROMEDIO DE NOTAS.": safe_str(row.get("promedio_pregrado")),
        "NOTA FINAL DE LICENCIATURA O TITULO PROFESIONAL O EQUIVALENTE.": safe_str(row.get("nota_final")),
        # Ranking
        "Posición de egreso del/de la estudiante al momento de finalizar su pregrado*:": safe_str(row.get("posicion_egreso")),
        "Total de estudiantes de su generación de egreso o titulación*.": safe_str(row.get("total_generacion")),
        "Ranking de egreso de pregrado, respecto de la generación de egreso o titulación": safe_str(row.get("ranking_porcentaje")),
    }
    replaced = replace_line_startswith(doc, label_map)
    missing = set(label_map.keys()) - replaced
    if missing:
        print(f"[Form1][WARN] No se pudieron ubicar {len(missing)} etiquetas, revise plantilla o textos:\n  - " + "\n  - ".join(missing))

def fill_form_2(doc, row):
    """
    Formulario N° 2: Certificado de Estado del/de la Postulante
    """
    nombre_completo = full_name(row)
    # Emulate checkboxes first
    mark_status_checkboxes(doc, row.get("estado_postulacion"))

    label_map = {
        "Nombre del postulante": nombre_completo,
        "Rut o número de pasaporte del postulante": safe_str(row.get("run") or row.get("pasaporte")),
        "Programa de destino (nombre del programa, según registro CNA-Chile)*": safe_str(row.get("programa_destino")),
        "Mención (si aplica, según registro CNA-Chile) *": safe_str(row.get("mencion_destino")),
        "Universidad (en caso de ser un programa en consorcio, señalar todas las universidades que lo integran) *": safe_str(row.get("universidad_destino")),
        "Región de los estudios de postgrado": safe_str(row.get("region_postgrado")),
        "Fecha de inicio de estudios (mes o semestre, año) *": safe_str(row.get("fecha_inicio_postgrado")),
        "**Nombre, cargo y firma de Autoridad Competente": f"{safe_str(row.get('autoridad_nombre'))}, {safe_str(row.get('autoridad_cargo'))}".strip(", "),
    }
    replaced = replace_line_startswith(doc, label_map)
    missing = set(label_map.keys()) - replaced
    if missing:
        print(f"[Form2][WARN] No se pudieron ubicar {len(missing)} etiquetas, revise plantilla o textos:\n  - " + "\n  - ".join(missing))

# -------- Main --------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--csv", default="estudiantes.csv", help="Ruta al CSV con datos de estudiantes")
    ap.add_argument("--f1", default="Formulario_N1_2026 (1).docx", help="Plantilla Formulario N°1 (DOCX)")
    ap.add_argument("--f2", default="Formulario_N_2_2026.docx", help="Plantilla Formulario N°2 (DOCX)")
    ap.add_argument("--outdir", default="salida", help="Directorio de salida")
    args = ap.parse_args()

    csv_path = Path(args.csv)
    f1_path = Path(args.f1)
    f2_path = Path(args.f2)
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    if not csv_path.exists():
        print(f"[ERROR] No existe el CSV: {csv_path}")
        sys.exit(1)
    if not f1_path.exists():
        print(f"[ERROR] No se encontró la plantilla F1: {f1_path}")
        sys.exit(1)
    if not f2_path.exists():
        print(f"[ERROR] No se encontró la plantilla F2: {f2_path}")
        sys.exit(1)

    df = pd.read_csv(csv_path)
    required = ["nombre", "apellido"]
    for col in required:
        if col not in df.columns:
            print(f"[ERROR] Falta columna obligatoria en CSV: {col}")
            sys.exit(1)

    for idx, row in df.iterrows():
        # Load fresh templates for each student
        doc1 = Document(str(f1_path))
        doc2 = Document(str(f2_path))

        fill_form_1(doc1, row)
        fill_form_2(doc2, row)

        apellido = safe_str(row.get("apellido")).replace(" ", "_")
        nombre = safe_str(row.get("nombre")).replace(" ", "_")
        base = f"{apellido}_{nombre}" if (apellido or nombre) else f"estudiante_{idx+1}"

        out1 = outdir / f"{base}_F1.docx"
        out2 = outdir / f"{base}_F2.docx"
        doc1.save(str(out1))
        doc2.save(str(out2))
        print(f"[OK] Generados: {out1.name} y {out2.name}")

if __name__ == "__main__":
    main()
